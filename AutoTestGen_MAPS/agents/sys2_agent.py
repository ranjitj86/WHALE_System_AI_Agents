from typing import Dict, List, Any, Union
from .base_agent import BaseAgent
import spacy
from transformers import pipeline
import pandas as pd
import json
import re
import uuid
import os
import io
import csv
from docx import Document
from fpdf import FPDF

class Sys2Agent(BaseAgent):
    """Agent 2: SYS.2 Requirement Drafting and Structuring"""
    
    def __init__(self, config: Dict[str, Any] = None):
        super().__init__(config)
        # Initialize NLP components - consider if a more powerful model is needed later
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except:
            print("Downloading en_core_web_sm model...")
            spacy.cli.download("en_core_web_sm")
            self.nlp = spacy.load("en_core_web_sm")
        
        # Initialize text classification pipeline - can be fine-tuned later
        try:
            self.classifier = pipeline("text-classification")
        except Exception as e:
            print(f"Could not load text classification pipeline: {e}")
            self.classifier = None # Handle case where pipeline fails to load

        self.templates = self._load_templates()
        self.setup_pipelines()
    
    def setup_pipelines(self):
        """Initialize custom NLP pipelines or models if needed"""
        # Example: Add a custom component for requirement structure analysis
        # if 'requirement_structure_analyzer' not in self.nlp.pipe_names:
        #     requirement_structure_analyzer = RequirementStructureAnalyzer(self.nlp)
        #     self.nlp.add_pipe("requirement_structure_analyzer", last=True)
        pass
    
    def validate(self, data: Any) -> bool:
        """Validate input data for Sys2Agent."""
        # Added comment to trigger reload
        if not isinstance(data, dict):
            return False
        
        source = data.get('source')
        
        if source == 'upload' or source == 'automatic':
            # For file sources, need 'file_path'
            return 'file_path' in data and isinstance(data['file_path'], str)
        elif source == 'raw_text':
            # For raw text, need 'raw_content'
            return 'raw_content' in data and isinstance(data['raw_content'], str)
        else:
            # Unknown source or missing source
            return False

    # Add basic logging methods to resolve the error temporarily
    def log_info(self, message: str):
        """Logs an informational message."""
        print(f"[INFO - Sys2Agent] {message}")

    def log_warning(self, message: str):
        """Logs a warning message."""
        print(f"[WARNING - Sys2Agent] {message}")

    def log_error(self, exception: Exception, message: str = ""):
        """Logs an error message and exception."""
        print(f"[ERROR - Sys2Agent] {message}: {exception}")

    def process_sys1_input(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Processes SYS.1 input from various sources (file, raw text) to draft and structure SYS.2 requirements.

        Args:
            input_data (Dict[str, Any]): A dictionary containing the input data.
                                         Expected keys:
                                         - 'source' (str): 'upload', 'automatic', or 'raw_text'
                                         - 'file_path' (str): Path to the file (if source is 'upload' or 'automatic')
                                         - 'raw_content' (str): Raw text content (if source is 'raw_text')

        Returns:
            Dict[str, Any]: A dictionary containing the processing results:
                            - 'status' (str): 'success' or 'error'
                            - 'message' (str): A message describing the result or error
                            - 'sys2_requirements' (List[Dict[str, Any]]): Drafted and structured SYS.2 requirements
                            - 'dependencies' (List[Dict[str, str]]): Identified dependencies
                            - 'classification' (Dict[str, str]): Classification results for SYS.2 requirements
                            - 'verification_mapping' (Dict[str, str]): Verification mapping for SYS.2 requirements
        """
        self.log_info(f"Starting Sys2Agent processing from source: {input_data.get('source')}")
        try:
            sys1_requirements = self._read_sys1_input(input_data)
            self.log_info(f"Read {len(sys1_requirements)} SYS.1 requirements.")

            sys2_requirements = []
            dependencies = []
            classification_results = {}
            verification_mapping_results = {}

            # Initialize a counter for sequential SYS.2 IDs
            sys2_counter = 1

            # Process each SYS.1 requirement
            for i, sys1_req in enumerate(sys1_requirements):
                # Ensure sys1_id and sys1_requirement keys exist
                sys1_id = sys1_req.get('SYS.1 Req. ID', f'SYS.1.AutoGen_{i+1}') # Use correct key
                original_sys1_text = sys1_req.get('SYS.1 System Requirement', '') # Get original SYS.1 text

                if not original_sys1_text.strip():
                    self.log_warning(f"Skipping empty or whitespace SYS.1 requirement for ID: {sys1_id}")
                    continue # Skip empty requirements

                # 1. Transform and Draft SYS.2 Requirement
                # Generate sequential SYS.2 ID (SYS.2-001, SYS.2-002, ...)
                sys2_id = f'SYS.2-{sys2_counter:03d}' # Format with leading zeros
                sys2_counter += 1 # Increment counter for the next requirement

                # Technically evaluate and rewrite the SYS.1 requirement
                drafted_sys2_text = self._technically_evaluate_and_rewrite(original_sys1_text)

                # 2. Auto-Fill Metadata
                priority = self._assign_priority(original_sys1_text) # Infer priority from text
                # Generate rationale with more context
                rationale = self._generate_rationale(
                    sys1_id,
                    original_sys1_text,
                    drafted_sys2_text, # Pass drafted text
                    self._determine_classification(drafted_sys2_text),    # Pass classification
                    self._map_verification(sys2_id, drafted_sys2_text) # Pass verification mapping
                ) # Link to SYS.1 source
                req_type = self._determine_type(original_sys1_text) # Determine type (Functional/Non-functional)
                domain = self._infer_domain(original_sys1_text) # Infer domain (Placeholder)
                release_planning = 'TBD' # Default release planning

                # 3. Classification
                # Assign classification as 'Functional' for all SYS.2 requirements as requested
                classification = "Functional"
                classification_results[sys2_id] = classification

                # 4. Dependency Mapping (Basic Placeholder)
                req_dependencies = self._map_single_requirement_dependencies(sys2_id, drafted_sys2_text)
                dependencies.extend(req_dependencies)

                # 5. Modularization (Placeholder)
                # If a requirement is complex, this is where you'd split it.
                # For now, we assume a 1:1 or 1:Many (if splitting) from SYS.1 to SYS.2.

                # 6. Custom Rule Engine (Placeholder)
                # Apply domain-specific rules here if implemented.
                # Note: apply_custom_rules should likely be applied after all initial drafting is done
                # For now, leaving it as is, but this might need refactoring.
                # sys2_requirements = self.apply_custom_rules(sys2_requirements) # Apply rules after appending?

                # 7. Verification Mapping
                # Assign verification mapping as 'System Qualification Test (SYS.5)' for all SYS.2 requirements as requested
                verification_mapping = "System Qualification Test (SYS.5)"
                verification_mapping_results[sys2_id] = verification_mapping

                # 8. Generate Verification Criteria
                verification_criteria = self._generate_verification_criteria(drafted_sys2_text, verification_mapping)

                sys2_requirements.append({
                    'sys1_id': sys1_id, # Keep track of the source SYS.1 ID
                    'sys1_requirement': original_sys1_text, # Store original SYS.1 text
                    'sys2_id': sys2_id,
                    'sys2_requirement': drafted_sys2_text,
                    'domain': domain,
                    'priority': priority,
                    'release_planning': release_planning,
                    'rationale': rationale,
                    'type': req_type, # This might be redundant with classification, decide on schema
                    'classification': classification,
                    'verification_mapping': verification_mapping,
                    'verification_criteria': verification_criteria, # Store generated criteria
                    'req_status': 'Draft' # Default status
                })

            # Apply custom rules to the generated SYS.2 requirements (moved outside the loop)
            sys2_requirements = self.apply_custom_rules(sys2_requirements)

            # Refine dependencies (Optional: further process the collected dependencies)
            # refined_dependencies = self._refine_dependencies(dependencies)

            self.log_info(f"Generated {len(sys2_requirements)} SYS.2 requirements.")
            return {
                'status': 'success',
                'message': 'SYS.2 requirements drafted and structured.',
                'sys2_requirements': sys2_requirements,
                'dependencies': dependencies,
                'classification': classification_results,
                'verification_mapping': verification_mapping_results
            }

        except FileNotFoundError as e:
            self.log_error(e, f"Input file not found: {e}")
            return {'status': 'error', 'message': f"Input file not found: {e}"}
        except Exception as e:
            self.log_error(e, "Error during Sys2Agent processing")
            return {'status': 'error', 'message': str(e)}
    
    def _read_sys1_input(self, input_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Reads SYS.1 requirements from the specified input source."""
        source = input_data.get('source')
        sys1_requirements = []

        if source in ['upload', 'automatic']:
            file_path = input_data.get('file_path')
            if not file_path or not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found at {file_path}")

            try:
                # Assuming Excel file with columns 'sys1_id' and 'sys1_requirement'
                df = pd.read_excel(file_path)
                self.log_info(f"Successfully read Excel file: {file_path}")
                self.log_info(f"DataFrame shape: {df.shape}")
                self.log_info(f"DataFrame columns: {df.columns.tolist()}")
                # Convert DataFrame rows to a list of dictionaries, ensuring None is handled
                sys1_requirements = df.where(pd.notnull(df), None).to_dict(orient='records')
                
                self.log_info(f"Raw requirements read from Excel (before cleaning): {sys1_requirements}")

                # Ensure required keys exist and are not None, defaulting to empty string if needed
                for req in sys1_requirements:
                    # Use .get() with a default empty string to prevent None
                    req['SYS.1 Req. ID'] = str(req.get('SYS.1 Req. ID', '')) if req.get('SYS.1 Req. ID') is not None else ''
                    req['SYS.1 System Requirement'] = str(req.get('SYS.1 System Requirement', '')) if req.get('SYS.1 System Requirement') is not None else ''

            except Exception as e:
                self.log_error(e, f"Error reading Excel file {file_path}")
                raise IOError(f"Error reading Excel file {file_path}: {e}")

        elif source == 'raw_text':
            raw_content = input_data.get('raw_content', '')
            if raw_content.strip():
                # Simple split by lines for raw text - could be more sophisticated
                lines = raw_content.strip().split('\n')
                for i, line in enumerate(lines):
                    if line.strip(): # Avoid adding empty lines
                        sys1_requirements.append({
                            'SYS.1 Req. ID': f'SYS.1.Raw_{i+1}',
                            'SYS.1 System Requirement': line.strip()
                        })
                self.log_info(f"Requirements read from raw text: {sys1_requirements}")

        # Filter out any entries where the requirement text is empty after reading
        # Ensure req.get() provides a default '' before calling strip()
        initial_count = len(sys1_requirements)
        sys1_requirements = [req for req in sys1_requirements if req.get('SYS.1 System Requirement', '').strip()]
        filtered_count = len(sys1_requirements)
        self.log_info(f"Filtered SYS.1 requirements: {filtered_count} remaining out of {initial_count}.")
        self.log_info(f"Filtered requirements: {sys1_requirements}")

        return sys1_requirements
    
    def _load_templates(self) -> Dict[str, Any]:
        """Load IREB and IEEE 830 compliant templates"""
        # Define templates - can be loaded from a config file later
        return {
            'functional': "The system shall {requirement}",
            'non_functional': "The system shall {requirement}",
            'assumption': "It is assumed that {requirement}",
            'constraint': "The system is subject to the constraint that {requirement}",
            # Add more templates as needed (e.g., performance, security)
        }
    
    def _apply_single_template(self, requirement_text: str, req_type: str) -> str:
        """Applies a specific template to a requirement text."""
        template = self.templates.get(req_type.lower(), "{requirement}") # Default to just requirement if type not found
        # Simple formatting - more complex template engines could be used
        try:
            return template.format(requirement=requirement_text)
        except KeyError:
            self.log_warning(f"Template for type '{req_type}' not found. Using raw text.")
            return requirement_text
    
    def _technically_evaluate_and_rewrite(self, sys1_requirement_text: str) -> str:
        """
        Performs a basic technical evaluation and rewrites the SYS.1 requirement.
        This is a placeholder for more sophisticated NLP/rule-based logic.
        """
        self.log_info(f"Technically evaluating and rewriting: {sys1_requirement_text[:80]}...")

        doc = self.nlp(sys1_requirement_text)
        evaluation_findings = []

        # Evaluation 1: Check for passive voice
        # Basic pattern: aux + verb_past + by (simplified)
        passive_voice_found = False
        for token in doc:
            if token.dep_ == 'nsubjpass':
                 # Found a passive subject, check for passive auxiliary
                 if any(child.dep_ == 'auxpass' for child in token.head.children):
                     passive_voice_found = True
                     break
        if passive_voice_found:
             evaluation_findings.append("Potential passive voice detected.")

        # Evaluation 2: Check for potential ambiguity markers (simple keyword check)
        ambiguity_markers = ['may', 'can', 'might', 'possibly', 'usually', 'generally']
        if any(token.text.lower() in ambiguity_markers for token in doc):
             evaluation_findings.append("Potential ambiguity marker detected.")

        # Placeholder for other checks (e.g., nominalizations, conjunctions like 'and/or')
        # evaluation_findings.extend(self._check_nominalizations(doc))
        # evaluation_findings.extend(self._check_conjunctions(doc))

        # Determine requirement type for templating
        req_type = self._determine_type(sys1_requirement_text)

        # --- Enhanced Rewriting Logic ---
        rewritten_text_core = sys1_requirement_text.strip() # Start with the original text

        # Basic attempt to make passive sentences active (simplified)
        if passive_voice_found:
            # This is a very basic heuristic and might not work for complex sentences
            # A more robust solution would require deeper dependency parsing and sentence transformation
            try:
                # Find the passive subject and the verb
                passive_subject = None
                passive_verb = None
                for token in doc:
                    if token.dep_ == 'nsubjpass':
                        passive_subject = token
                        if token.head.pos_ == 'VERB':
                             passive_verb = token.head
                        break # Assume one main passive structure for simplicity
                
                if passive_subject and passive_verb:
                     # Attempt to restructure: [Optional Agent] [passive_verb] [passive_subject]
                     # This is highly simplified and likely needs domain-specific rules or advanced NLP
                     # For now, let's just note that passive voice was addressed.
                     rewritten_text_core = f"[Passive voice considered] {sys1_requirement_text.strip()}"
                else:
                    # If we can't identify subject/verb for restructuring, just note it
                    rewritten_text_core = f"[Passive voice detected] {sys1_requirement_text.strip()}"

            except Exception as e:
                self.log_warning(f"Error attempting passive voice rewrite: {e}")
                rewritten_text_core = f"[Passive voice detected] {sys1_requirement_text.strip()}"
        
        # Add notes about other findings if any (excluding passive voice, if handled in rewrite)
        other_findings = [f for f in evaluation_findings if "passive voice" not in f]
        if other_findings:
            rewritten_text_core += f" [Review: {'; '.join(other_findings)}]"

        # Apply the template to the potentially rewritten/noted core text
        final_sys2_text = self._apply_single_template(rewritten_text_core, req_type)

        # --- End Enhanced Rewriting Logic ---

        return final_sys2_text
    
    def _transform_to_sys2(self, sys1_requirements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Deprecated: Use process_sys1_input instead."""
        pass # This method is being replaced
    
    def _apply_templates(self, requirements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Deprecated: Template application is now part of process_sys1_input."""
        pass # This method is being replaced
    
    def _map_dependencies(self, requirements: List[Dict[str, Any]]) -> Dict[str, List[str]]:
        """Deprecated: Dependency mapping is now handled within process_sys1_input."""
        pass # This method is being replaced
    
    def _classify_requirements(self, requirements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Deprecated: Classification is now handled within process_sys1_input."""
        pass # This method is being replaced
    
    def _determine_type(self, content: str) -> str:
        """Determine requirement type (Functional/Non-functional) - Basic heuristic."""
        # Basic heuristic: look for keywords. Can be improved with NLP.
        content_lower = content.lower()
        if any(keyword in content_lower for keyword in ['shall', 'must', 'will']):
            return 'Functional'
        elif any(keyword in content_lower for keyword in ['should', 'may']):
            return 'Non-functional'
        elif any(keyword in content_lower for keyword in ['assume', 'assuming', 'assumption']):
            return 'Assumption'
        elif any(keyword in content_lower for keyword in ['constraint', 'limit', 'restrict']):
            return 'Constraint'
        else:
            return 'Other'
    
    def _assign_priority(self, requirement_text: str) -> str:
        """Assign priority based on simple keyword matching - Placeholder."""
        # More sophisticated priority assignment would be needed
        text_lower = requirement_text.lower()
        if "critical" in text_lower or "essential" in text_lower:
            return 'High'
        elif "important" in text_lower or "should" in text_lower:
            return 'Medium'
        else:
            return 'Low'
    
    def _generate_rationale(
        self,
        sys1_id: str,
        original_sys1_text: str,
        drafted_sys2_text: str,
        classification: str,
        verification_mapping: str
    ) -> str:
        """Generate rationale linking to the source SYS.1 ID and including analysis details."""
        rationale_parts = [
            f"Derived from SYS.1 requirement {sys1_id}.",
            f"Original text: {original_sys1_text}"
        ]

        # Check if any technical review notes were added during rewriting
        review_match = re.search(r' \[Technical Review: (.*?)\]', drafted_sys2_text)
        if review_match:
            review_notes = review_match.group(1)
            rationale_parts.append(f"Technical review noted: {review_notes}.")

        # Include classification and verification mapping
        rationale_parts.append(f"Classified as: {classification}.")
        rationale_parts.append(f"Verification method: {verification_mapping}.")

        # You could add more comparison logic here between original and drafted text
        # Example: if original_sys1_text != drafted_sys2_text:
        # rationale_parts.append("Agent 2 made modifications during drafting.")

        return " ".join(rationale_parts)
    
    def _determine_classification(self, requirement_text: str) -> str:
        """Classify requirement (Functional, Non-functional, Assumption, Constraint) using NLP pipeline or basic heuristic."""
        if self.classifier:
            try:
                # Use the text classification pipeline
                # The labels from the pipeline might not directly match our categories.
                # We'll need to interpret the pipeline's output.
                # For a generic pipeline, it might return sentiment or topics. This is a placeholder interpretation.
                classification_results = self.classifier(requirement_text)
                # Assuming the pipeline returns a list of dicts like [{'label': 'LABEL_XXX', 'score': Y.YY}]
                if classification_results and isinstance(classification_results, list) and classification_results[0]:
                    label = classification_results[0]['label'].lower()
                    # Simple mapping from potential generic labels to our categories
                    if 'functional' in label or 'system' in label:
                        return 'Functional'
                    elif 'non-functional' in label or 'performance' in label or 'security' in label or 'usability' in label:
                        return 'Non-functional'
                    elif 'assumption' in label or 'assume' in label:
                        return 'Assumption'
                    elif 'constraint' in label or 'restrict' in label or 'limit' in label:
                        return 'Constraint'
                    else:
                        return 'Other' # Default if classification is unclear
            except Exception as e:
                self.log_error(e, "Error during text classification pipeline execution")
                # Fallback to basic heuristic if pipeline fails
                return self._determine_type(requirement_text)
        else:
            # Fallback to basic heuristic if classifier is not loaded
            return self._determine_type(requirement_text)
    
    def _infer_domain(self, requirement_text: str) -> str:
        """Infer domain from requirement text - Placeholder."""
        # Requires domain-specific knowledge or NLP topic modeling.
        return 'Software' # Default domain changed to Software
    
    def _map_single_requirement_dependencies(self, sys2_id: str, requirement_text: str) -> List[Dict[str, str]]:
        """Identifies dependencies for a single requirement using NLP and keyword matching."""
        dependencies = []
        doc = self.nlp(requirement_text) # Process text with spaCy

        # 1. Look for mentions of other potential requirement IDs (e.g., SYS.1.X, SYS.2.Y)
        # This regex assumes IDs follow the pattern SYS.[1 or 2].<Alphanumeric+>
        potential_ids_in_text = re.findall(r'(SYS\.[12]\.\w+)', requirement_text)

        for dep_id in potential_ids_in_text:
            # Avoid creating a dependency on itself and ensure the mentioned ID looks valid
            if dep_id != sys2_id and re.match(r'SYS\.[12]\.\w+', dep_id):
                dependencies.append({
                    'from': sys2_id,
                    'to': dep_id,
                    'label': 'mentions' # Label indicating the dependency type
                })

        # 2. (Placeholder) Analyze sentence structure or keywords for implicit dependencies
        # This is a complex task requiring more advanced NLP rules or models.
        # Example: Look for causal verbs, conditional phrases, etc.
        # for sent in doc.sents:
        #     # Analyze sentence.text or sentence.root.text for dependency indicators
        #     # if 'depends on' in sent.text.lower():
        #     #     # Attempt to extract the dependency relationship and target
        #     #     pass

        # Remove duplicate dependencies
        # Convert list of dicts to a set of tuples for uniqueness, then back to list of dicts
        unique_dependencies = set()
        for dep in dependencies:
            # Convert dict to a tuple of (key, value) pairs for hashing
            unique_dependencies.add(tuple(sorted(dep.items())))

        return [dict(t) for t in unique_dependencies]
    
    def _map_verification(self, sys2_id: str, requirement_text: str) -> str:
        """Map requirement to verification method using basic keyword matching."""
        text_lower = requirement_text.lower()

        # Check for System Qualification Test keywords first
        if any(keyword in text_lower for keyword in ['system test', 'qualification test', 'acceptance test']):
            return 'System Qualification Test (SYS.5)'

        # Existing checks for other verification methods
        elif any(keyword in text_lower for keyword in ['test', 'verify', 'validate', 'check']):
            return 'Test'
        elif any(keyword in text_lower for keyword in ['analyze', 'analysis', 'study']):
            return 'Analysis'
        elif any(keyword in text_lower for keyword in ['inspect', 'inspection', 'review']):
            return 'Inspection'
        else:
            return 'Test (Default)' # Default verification method

    def _generate_verification_criteria(self, sys2_requirement_text: str, verification_method: str) -> str:
        """
        Generates basic verification criteria based on requirement text and method.
        This is a placeholder and can be enhanced with more sophisticated logic.
        """
        self.log_info(f"Generating verification criteria for: {sys2_requirement_text[:80]}...")
        doc = self.nlp(sys2_requirement_text)

        # Basic NLP analysis to extract potential key elements
        main_verb = ''
        main_subject = ''
        main_object = ''
        conditions = []

        for token in doc:
            if token.dep_ == 'ROOT' and token.pos_ == 'VERB':
                main_verb = token.text
                main_subject = next((child.text for child in token.children if child.dep_ == 'nsubj' or child.dep_ == 'nsubjpass'), '')
                main_object = next((child.text for child in token.children if child.dep_ == 'dobj'), '')
            # Simple way to find clauses that might contain conditions (e.g., starting with 'if', 'when')
            if token.dep_ == 'advcl' or token.dep_ == 'acl' or token.text.lower() in ['if', 'when', 'where']:
                 # Attempt to capture the clause text
                 conditions.append(token.sent.text[token.i:].strip())

        # Construct criteria based on verification method and extracted elements
        criteria = "Verification criteria TBD."

        if verification_method == 'System Qualification Test (SYS.5)' or verification_method == 'Test' or verification_method == 'Test (Default)':
            criteria = f"Develop test case(s) to verify that {main_subject} can {main_verb} {main_object} "
            if conditions:
                criteria += f" when/if { ' and '.join(conditions) }."
            else:
                 criteria += "."
             # Add a general check
            criteria += f" Specifically, verify compliance with: {sys2_requirement_text.strip()}"

        elif verification_method == 'Analysis':
            criteria = f"Perform analysis to confirm that {main_subject} can {main_verb} {main_object}. "
            if conditions:
                criteria += f" The analysis should consider conditions such as: { ' and '.join(conditions) }."
            criteria += f" Review documentation/models related to: {sys2_requirement_text.strip()}"

        elif verification_method == 'Inspection':
             criteria = f"Inspect documentation and implementation to ensure {main_subject} can {main_verb} {main_object}. "
             if conditions:
                 criteria += f" Pay close attention to aspects related to: { ' and '.join(conditions) }."
             criteria += f" Specifically, check for: {sys2_requirement_text.strip()}"

        return criteria.strip()

    # Deprecated method, replaced by process_sys1_input
    def process(self, input_data: Any) -> Dict[str, Any]:
        """Deprecated: Use process_sys1_input instead."""
        self.log_warning("Calling deprecated process method. Use process_sys1_input instead.")
        return {'status': 'error', 'message': 'Deprecated method.'}

    # Deprecated method, replaced by process_sys1_input
    def process_sys1_requirements(self, sys1_requirements):
        """Deprecated: Use process_sys1_input instead."""
        self.log_warning("Calling deprecated process_sys1_requirements method. Use process_sys1_input instead.")
        return [], [] # Return empty lists for deprecated method

    # Placeholder method for future analysis/structuring logic
    def analyze_dependencies(self, requirements):
        """Placeholder for more advanced dependency analysis."""
        pass

    # Placeholder method for future refinement logic
    def refine_requirements(self, requirements):
        """Placeholder for requirement refinement based on feedback or rules."""
        pass

    # Placeholder for integrating Agent 3 feedback
    def integrate_agent3_feedback(self, sys2_requirements: List[Dict[str, Any]], feedback: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Placeholder for integrating feedback from Agent 3."""
        # Logic to update sys2_requirements based on feedback
        return sys2_requirements

    # Placeholder for applying custom rules
    def apply_custom_rules(self, sys2_requirements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Placeholder for applying domain-specific custom rules."""
        # Logic to modify requirements based on custom rules
        self.log_info("Applying custom rules (placeholder).")
        # Example of how custom rules might modify a requirement:
        # for req in sys2_requirements:
        #     if "performance" in req.get('sys2_requirement', '').lower():
        #         req['priority'] = 'High'
        return sys2_requirements

    def export_requirements(self, requirements: List[Dict[str, Any]], format: str, export_fields_list: List[str] | None = None) -> Union[io.BytesIO, str, None]:
        """Exports SYS.2 requirements to the specified format.

        Args:
            requirements (List[Dict[str, Any]]): The list of SYS.2 requirements.
            format (str): The desired export format (e.g., 'xlsx', 'csv', 'docx', 'pdf', 'txt').
            export_fields_list (List[str] | None): An optional list of field keys to include in the export. If None, defaults are used.

        Returns:
            Union[io.BytesIO, str, None]: The exported data in the specified format, or None if format is unsupported.
                                        For binary formats (xlsx, docx, pdf), returns io.BytesIO.
                                        For text formats (csv, txt), returns str.

        Raises:
            ValueError: If the format is unsupported.
        """
        if not requirements:
            return None # Or raise an error, depending on desired behavior for empty data

        # Define the columns/fields to include in the export
        # Ensure the order is consistent with the requested format
        # Use the provided export_fields_list if available, otherwise use default
        fields_to_export = export_fields_list if export_fields_list is not None else [
            'sys1_id', 'sys1_requirement', # Corresponding to "SYS.1 Req. ID", "SYS.1 Requirement"
            'sys2_id', 'sys2_requirement', # Corresponding to "SYS.2 Req. ID", "SYS.2 System Requirement"
            'classification', # Corresponding to "Type (Functional, Non-Functional)"
            'verification_mapping', # Corresponding to "Verification Method (SYS.5)"
            'verification_criteria', # Corresponding to "Verification Criteria" (Placeholder)
            'domain', 'priority', 'rationale', 'req_status' # Remaining fields
            # 'Actions' is a UI element and not included in data export
        ]

        # Define the header mapping from internal keys to requested headers
        # Only include headers for the fields being exported
        header_mapping = {
            'sys1_id': 'SYS.1 Req. ID',
            'sys1_requirement': 'SYS.1 Requirement',
            'sys2_id': 'SYS.2 Req. ID',
            'sys2_requirement': 'SYS.2 System Requirement',
            'classification': 'Type (Functional, Non-Functional)',
            'verification_mapping': 'Verification Method (SYS.5)',
            'verification_criteria': 'Verification Criteria',
            'domain': 'Domain',
            'priority': 'Priority',
            'rationale': 'Rationale',
            'req_status': 'Requirement Status'
        }

        # Prepare data in a consistent format for export using the defined fields and headers
        export_data = []
        for req in requirements:
            row = {}
            for field in fields_to_export:
                # Use the header_mapping to get the desired column header
                header = header_mapping.get(field, field) # Default to field name if not in mapping
                # Get the data from the requirement dictionary, defaulting to empty string
                row[header] = req.get(field, '')
            export_data.append(row)

        if format == 'xlsx':
            print("[DEBUG] Exporting SYS.2 requirements:", export_data)  # Debug print
            df = pd.DataFrame(export_data)
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                df.to_excel(writer, index=False, sheet_name='SYS2 Requirements')
            output.seek(0)

            # --- Save to Inputs directory on the server ---
            try:
                output_path = r'D:\AgentX\AutoTestGen_MAPS_Agents123\AutoTestGen_MAPS\Inputs\sys2_requirements.xlsx'
                os.makedirs(os.path.dirname(output_path), exist_ok=True)
                df.to_excel(output_path, index=False)
                print(f"[DEBUG] Saved XLSX to {output_path}")
            except Exception as e:
                print(f"[ERROR] Could not save XLSX to Inputs directory: {e}")
            # ------------------------------------------------

            return output

        elif format == 'csv':
            output = io.StringIO()
            # Use the headers from the export_data (which respects fields_to_export) as fieldnames for CSV header
            fieldnames = export_data[0].keys() if export_data else []
            writer = csv.DictWriter(output, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(export_data)
            return output.getvalue()

        elif format == 'docx':
            document = Document()
            document.add_heading('SYS.2 Requirements', 0)

            # Add a table
            # Use the number of fields in fields_to_export for column count
            table = document.add_table(rows=1, cols=len(fields_to_export))
            table.style = 'Grid Table 4' # Apply a table style

            # Add header row using the header_mapping values for the selected fields
            header_cells = table.rows[0].cells
            for i, field in enumerate(fields_to_export):
                 header_cells[i].text = header_mapping.get(field, field) # Use mapped header

            # Add data rows
            for row_data in export_data:
                cells = table.add_row().cells
                # Iterate through fields_to_export to get data in correct order
                for i, field in enumerate(fields_to_export):
                     # Access data using the mapped header key from export_data
                     cells[i].text = str(row_data.get(header_mapping.get(field, field), ''))

            output = io.BytesIO()
            document.save(output)
            output.seek(0);
            return output

        elif format == 'pdf':
            class PDF(FPDF):
                def header(self):
                    self.set_font('Arial', 'B', 12)
                    self.cell(0, 10, 'SYS.2 Requirements Report', 0, 1, 'C')
                    self.ln(10)

                def footer(self):
                    self.set_y(-15)
                    self.set_font('Arial', 'I', 8)
                    self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', 0, 0, 'C')

            pdf = PDF()
            pdf.alias_nb_pages()
            # Determine orientation based on number of columns, or use Landscape by default for potentially wide tables
            orientation = 'L' if len(fields_to_export) > 5 else 'P' # Simple heuristic for orientation
            pdf.add_page(orientation=orientation)
            pdf.set_auto_page_break(auto=True, margin=15)

            # Define column widths - This will be more complex with variable fields.
            # For simplicity, we can try to distribute available width, or use fixed widths for known fields.
            # A better approach for a general export would be dynamic width calculation or configuration.
            # For this specific request (SYS.2 ID and SYS.2 Requirement), we can use fixed widths.
            if export_fields_list == ['sys2_id', 'sys2_requirement']:
                 col_widths = [40, 150] # Example widths for just these two columns
            elif len(fields_to_export) == 11: # Assume default 11 columns
                 col_widths = [
                    25, 45, 25, 45, 25, 25, 35, 20, 15, 30, 25 # Default widths for 11 columns
                 ]
            else:
                # Fallback: simple equal distribution of width
                page_width = pdf.w - 2 * pdf.l_margin
                col_width = page_width / len(fields_to_export)
                col_widths = [col_width] * len(fields_to_export)

            # Add table headers using the header_mapping values for the selected fields
            pdf.set_font('Arial', 'B', 8) # Smaller font for more columns
            for i, field in enumerate(fields_to_export):
                 # Use multi_cell for headers that might wrap
                 pdf.multi_cell(col_widths[i], 10, header_mapping.get(field, field), 1, 'C', False, 0)
            pdf.ln()

            # Add table rows
            pdf.set_font('Arial', '', 7) # Even smaller font for data
            # Truncation is less critical if only two fields are exported, but keeping the logic
            truncate_limit = 150 # Define truncation limit
            for row_data in export_data:
                for i, field in enumerate(fields_to_export):
                    header = header_mapping.get(field, field)
                    cell_data = str(row_data.get(header, '')) # Get data as string
                    
                    # Apply truncation to specific columns if they are among the exported fields
                    if field in ['verification_criteria', 'rationale'] and len(cell_data) > truncate_limit:
                        cell_data = cell_data[:truncate_limit] + '...'

                    # Use multi_cell for data that might wrap
                    pdf.multi_cell(col_widths[i], 10, cell_data, 1, 'L', False, 0)
                pdf.ln()

            output = io.BytesIO()
            # Use dest='S' to get the raw bytes output
            pdf_bytes = pdf.output(dest='S').encode('latin-1') # Output as string first, then encode
            output.write(pdf_bytes)
            output.seek(0);
            return output

        elif format == 'txt':
            output = io.StringIO()
            # Use the headers from export_data as fieldnames for TXT header
            output.write('	'.join([header for header in export_data[0].keys()]) + '\n' if export_data else '') # Ensure header is written correctly
            for row_data in export_data:
                 # Iterate through the keys (headers) of the row_data dictionary to maintain order and use mapped headers
                 output.write('	'.join([str(row_data.get(header, '')) for header in export_data[0].keys()]) + '\n')
            return output.getvalue()

        else:
            # Raise ValueError for unsupported formats
            raise ValueError(f"Unsupported export format: {format}")

    def get_dashboard_summary(self, requirements: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Calculates summary statistics for the SYS.2 dashboard.

        Args:
            requirements (List[Dict[str, Any]]): The list of SYS.2 requirements.

        Returns:
            Dict[str, Any]: A dictionary containing summary counts.
        """
        self.log_info(f"Calculating dashboard summary for {len(requirements)} requirements.")

        total_sys2_reqs = len(requirements)

        # Traceability to SYS.1 (assuming sys1_id is present if traced)
        traced_to_sys1_count = sum(1 for req in requirements if req.get('sys1_id'))
        not_traced_to_sys1_count = total_sys2_reqs - traced_to_sys1_count

        # Breakdown by Status
        status_counts: Dict[str, int] = {'Draft': 0, 'Reviewed': 0, 'Approved': 0, 'Rejected': 0, 'Other': 0}
        for req in requirements:
            status = req.get('req_status', 'Other')
            status_counts[status] = status_counts.get(status, 0) + 1

        # Breakdown by Classification
        classification_counts: Dict[str, int] = {'Functional': 0, 'Non-Functional': 0, 'Assumption': 0, 'Constraint': 0, 'Other': 0}
        for req in requirements:
            classification = req.get('classification', 'Other')
            classification_counts[classification] = classification_counts.get(classification, 0) + 1

        # Breakdown by Verification Method
        verification_counts: Dict[str, int] = {
            'System Qualification Test (SYS.5)': 0,
            'Test': 0,
            'Analysis': 0,
            'Inspection': 0,
            'Test (Default)': 0,
            'Other': 0 # Catch any unexpected values
        }
        for req in requirements:
             method = req.get('verification_mapping', 'Other')
             verification_counts[method] = verification_counts.get(method, 0) + 1

        summary = {
            'total_sys2_reqs': total_sys2_reqs,
            'traced_to_sys1': traced_to_sys1_count,
            'not_traced_to_sys1': not_traced_to_sys1_count,
            'status_breakdown': status_counts,
            'classification_breakdown': classification_counts,
            'verification_breakdown': verification_counts
        }

        self.log_info(f"Dashboard summary calculated: {summary}")
        return summary 